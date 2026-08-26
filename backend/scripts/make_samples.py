"""Generate the three FICTIONAL sample documents into backend/sample_data/.

All companies, people, and figures are invented for testing Alpha Detective
(CLAUDE_CODE_PROMPT.md §7). Figures are exact per spec:

1. meridian_q2_fy2026_earnings_call.pdf  — ~2-page reportlab earnings-call
   transcript (operator, prepared remarks, Q&A) with a real rendered
   quarterly-metrics table.
2. northwind_retail_q2_2026_earnings.txt — plain-text earnings release.
3. helios_energy_fy2025_annual_report.docx — annual report with a docx table.

Run:  cd backend && .venv/bin/python scripts/make_samples.py
"""

from __future__ import annotations

from pathlib import Path

SAMPLE_DIR = Path(__file__).resolve().parent.parent / "sample_data"

FICTIONAL_NOTE = "Fictional sample document generated for Alpha Detective testing."


# --------------------------------------------------------------------------- #
# 1. Meridian Systems — PDF earnings-call transcript with a metrics table
# --------------------------------------------------------------------------- #

def make_meridian_pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleC", parent=styles["Title"], fontSize=16, spaceAfter=6)
    sub = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=10, textColor=colors.grey, spaceAfter=14)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=8)
    speaker = ParagraphStyle("Speaker", parent=body, spaceBefore=2)

    def line(who: str, text: str):
        return Paragraph(f"<b>{who}:</b> {text}", speaker)

    story = [
        Paragraph("Meridian Systems, Inc. (MRDN)", title),
        Paragraph(
            "Q2 FY2026 Earnings Conference Call Transcript — August 14, 2025, 5:00 PM ET. "
            + FICTIONAL_NOTE,
            sub,
        ),
        Paragraph("Call Participants", h2),
        Paragraph(
            "Daniel Okafor — Chief Executive Officer; Priya Raghavan — Chief Financial "
            "Officer; Lena Hartwell — VP, Investor Relations.",
            body,
        ),
        Paragraph("Operator", h2),
        line(
            "Operator",
            "Good afternoon, and welcome to the Meridian Systems second quarter fiscal "
            "2026 earnings conference call. All participants are in listen-only mode. "
            "After the prepared remarks there will be a question-and-answer session. I "
            "would now like to turn the call over to Lena Hartwell, VP of Investor "
            "Relations. Please go ahead.",
        ),
        line(
            "Lena Hartwell",
            "Thank you, operator. During today's call we will make forward-looking "
            "statements and refer to non-GAAP measures. With me today are our CEO, "
            "Daniel Okafor, and our CFO, Priya Raghavan. Daniel, over to you.",
        ),
        Paragraph("Prepared Remarks", h2),
        line(
            "Daniel Okafor (CEO)",
            "Thank you, Lena, and good afternoon, everyone. Meridian delivered an "
            "outstanding second quarter. Revenue for the second quarter was $48.2 "
            "million, an increase of 23% year-over-year, driven by continued strength "
            "in our enterprise workflow platform. Annual recurring revenue reached "
            "$210.4 million, and net revenue retention was 118%, reflecting healthy "
            "expansion within our installed base. Given the momentum in the business, "
            "we are raising our full-year FY2026 revenue guidance to $196–200 million.",
        ),
        line(
            "Priya Raghavan (CFO)",
            "Thanks, Daniel. Turning to profitability and the balance sheet. Non-GAAP "
            "operating margin for the second quarter was 11%, up from 6% a year ago as "
            "we continued to scale efficiently. GAAP net loss was $(3.1) million, "
            "which includes stock-based compensation expense. We ended the quarter "
            "with $312 million in cash, cash equivalents, and short-term investments, "
            "and no debt. Headcount at quarter end was 1,240 employees. The table of "
            "quarterly metrics on the following page summarizes the quarter.",
        ),
        PageBreak(),
        Paragraph("Quarterly Metrics — Q2 FY2026", h2),
        Table(
            [
                ["Metric", "Q2 FY2026", "Commentary"],
                ["Revenue", "$48.2 million", "+23% year-over-year"],
                ["Annual recurring revenue (ARR)", "$210.4 million", "quarter end"],
                ["Net revenue retention (NRR)", "118%", "trailing twelve months"],
                ["Non-GAAP operating margin", "11%", "up from 6% a year ago"],
                ["GAAP net loss", "$(3.1) million", "includes SBC expense"],
                ["Cash and short-term investments", "$312 million", "no debt"],
                ["Headcount", "1,240", "quarter end"],
                ["FY2026 revenue guidance", "$196–200 million", "raised"],
            ],
            colWidths=[2.4 * inch, 1.6 * inch, 2.4 * inch],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            ),
        ),
        Spacer(1, 14),
        Paragraph("Question-and-Answer Session", h2),
        line(
            "Operator",
            "We will now begin the question-and-answer session. Our first question "
            "comes from Marcus Bell of Halcyon Research.",
        ),
        line(
            "Marcus Bell (Halcyon Research)",
            "Congratulations on the quarter. Daniel, can you unpack what is driving "
            "the raised FY2026 guidance of $196–200 million — is it new logos or "
            "expansion?",
        ),
        line(
            "Daniel Okafor (CEO)",
            "Thanks, Marcus. It is both, but expansion is the larger driver — net "
            "revenue retention of 118% tells that story. Our largest customers are "
            "standardizing more teams on the platform.",
        ),
        line(
            "Sofia Grant (Beacon Capital Markets)",
            "Priya, how should we think about the durability of the 11% non-GAAP "
            "operating margin while you keep hiring?",
        ),
        line(
            "Priya Raghavan (CFO)",
            "We ended the quarter at 1,240 employees and plan measured hiring in "
            "go-to-market. With $312 million of cash and no debt we can invest while "
            "expanding margin modestly for the full year.",
        ),
        line(
            "Operator",
            "This concludes today's question-and-answer session and the Meridian "
            "Systems Q2 FY2026 earnings call. Thank you for joining. You may now "
            "disconnect.",
        ),
    ]

    SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Meridian Systems Q2 FY2026 Earnings Call",
    ).build(story)


# --------------------------------------------------------------------------- #
# 2. Northwind Retail Group — plain-text earnings release
# --------------------------------------------------------------------------- #

NORTHWIND_TXT = f"""Northwind Retail Group, Inc. — Second Quarter 2026 Earnings Release
{FICTIONAL_NOTE}

SEATTLE — Northwind Retail Group, Inc. today reported financial results for its
second quarter of fiscal 2026.

Second Quarter 2026 Highlights

- Total revenue was $1.84 billion, an increase of 4.1% compared with the prior-year
  quarter.
- Comparable same-store sales increased 2.6%, led by grocery and home categories.
- E-commerce revenue grew 18% year-over-year and represented a record share of
  total sales.
- Gross margin was 33.9%, an improvement of 70 basis points versus the prior-year
  quarter, driven by lower freight costs and disciplined promotions.
- Diluted earnings per share were $1.12, compared with $0.98 in the second quarter
  of fiscal 2025.
- The company operated 214 stores at quarter end.
- The board of directors declared a quarterly cash dividend of $0.32 per share.

Chief Executive Officer Commentary

"Our teams delivered a strong quarter in a choppy consumer environment," said the
company. "Same-store sales growth of 2.6% and e-commerce growth of 18% show that
our omnichannel investments are paying off, while 70 basis points of gross margin
expansion to 33.9% demonstrates disciplined execution."

Capital Allocation

Northwind returned cash to shareholders through its quarterly dividend of $0.32
per share and continued its store refresh program across the 214-store fleet.

About Northwind Retail Group

Northwind Retail Group, Inc. is a fictional omnichannel retailer operating
general-merchandise stores and a growing e-commerce business. This document was
generated as sample data and describes no real company.
"""


def make_northwind_txt(path: Path) -> None:
    path.write_text(NORTHWIND_TXT, encoding="utf-8")


# --------------------------------------------------------------------------- #
# 3. Helios Energy plc — DOCX annual report with a table
# --------------------------------------------------------------------------- #

def make_helios_docx(path: Path) -> None:
    import docx

    d = docx.Document()
    d.add_heading("Helios Energy plc — Annual Report FY2025", level=0)
    d.add_paragraph(FICTIONAL_NOTE)

    d.add_heading("Chairman's Statement", level=1)
    d.add_paragraph(
        "Fiscal 2025 was a year of disciplined delivery for Helios Energy plc. Group "
        "revenue was $6.3 billion, and adjusted EBITDA reached $1.9 billion, "
        "reflecting resilient generation output and tight cost control. We ended the "
        "year with net debt of $4.1 billion, comfortably within our target leverage "
        "range."
    )

    d.add_heading("Strategic Progress", level=1)
    d.add_paragraph(
        "Our renewables portfolio reached 3.2 GW of installed capacity, anchored by "
        "the commissioning of the Solara Ridge and North Moor wind projects. Helios "
        "employed approximately 8,500 people at year end across generation, networks, "
        "and corporate functions."
    )

    d.add_heading("Outlook and Capital Allocation", level=1)
    d.add_paragraph(
        "For FY2026 the board has set capital expenditure guidance of $1.1 billion, "
        "weighted toward grid modernization and new renewable capacity. The board "
        "reaffirmed the group's dividend policy of a 40% payout of adjusted earnings."
    )

    d.add_heading("Financial Summary", level=1)
    rows = [
        ("Metric", "FY2025"),
        ("Revenue", "$6.3 billion"),
        ("Adjusted EBITDA", "$1.9 billion"),
        ("Net debt", "$4.1 billion"),
        ("Installed renewable capacity", "3.2 GW"),
        ("Employees", "8,500"),
        ("FY2026 capex guidance", "$1.1 billion"),
        ("Dividend payout policy", "40% of adjusted earnings"),
    ]
    table = d.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (metric, value) in enumerate(rows):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value

    d.add_paragraph(
        "Helios Energy plc is a fictional integrated energy group created as sample "
        "data; it describes no real company."
    )
    d.save(str(path))


# --------------------------------------------------------------------------- #

def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    pdf = SAMPLE_DIR / "meridian_q2_fy2026_earnings_call.pdf"
    txt = SAMPLE_DIR / "northwind_retail_q2_2026_earnings.txt"
    dcx = SAMPLE_DIR / "helios_energy_fy2025_annual_report.docx"
    make_meridian_pdf(pdf)
    make_northwind_txt(txt)
    make_helios_docx(dcx)
    for p in (pdf, txt, dcx):
        print(f"wrote {p} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
