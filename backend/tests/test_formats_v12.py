"""New v1.2 ingest formats and extraction caps -- CONTRACTS.md SS1.3, SS2 ingest.py.

Every fixture in this file is generated in-test (tiny, deterministic, no golden
files checked in). The load-bearing case is `test_xlsx_numeric_cell_is_retrievable`:
an .xlsx whose numeric cell must come back through a REAL PROVIDER=none query.
A format that parses but is not retrievable is a format that does not work.

Caps (SS2) are named module-top constants precisely so they can be exercised
cheaply. Where generating a genuinely oversized artifact would cost seconds of
suite time (200k cells, 500 slides, 5M chars) the constant is patched down and
the frozen default is asserted separately by
`test_extraction_cap_constants_are_frozen`. The depth bomb, the sheet cap and
the zip bombs are exercised for real.

SS1.3 (law): hitting a cap is a clean per-file failure with HTTP 200 -- never a
500, never a crash, never a partially committed document.
"""

import json as _json
import zlib
from types import SimpleNamespace

import pytest

from conftest import (
    ALLOWED_EXTS_V12,
    EXTRACTION_CAPS,
    app_client,
    make_html,
    make_json_bytes,
    make_pptx,
    make_xlsx,
    make_zip_bytes,
    nested_json,
    post_query,
    upload_bytes,
)

GOOD_TXT = ("sentinel.txt", b"Sentinel document used to prove the batch still commits.\n")


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------
@pytest.fixture()
def fmt(tmp_path, samples, qa):
    """Fresh, empty keyless app -- one corpus per test keeps retrieval unambiguous."""
    storage = tmp_path / "storage"
    with app_client(storage) as client:
        yield SimpleNamespace(client=client, storage=storage, qa=qa)


def index_one(client, name, data):
    resp = upload_bytes(client, [(name, data)])
    assert resp.status_code == 200, f"HTTP {resp.status_code}: {resp.text[:400]}"
    entry = resp.json()["documents"][0]
    assert entry["status"] == "indexed", (
        f"{name} failed to index: {entry.get('error')!r} (full entry: {entry})"
    )
    return entry


def previews(client, doc_id):
    resp = client.get(f"/api/documents/{doc_id}/chunks")
    assert resp.status_code == 200, resp.text[:300]
    return resp.json()["chunks"]


def all_text(client, doc_id):
    return " ".join(r["preview"] for r in previews(client, doc_id))


def patch_cap(qa, name, value):
    return qa.patch_backend_attr(name, value)


def assert_clean_per_file_failure(f, bad_name, bad_bytes, expected_error):
    """SS1.3 (law): per-file failure, HTTP 200, batch survives, nothing persisted."""
    resp = upload_bytes(f.client, [(bad_name, bad_bytes), GOOD_TXT])
    assert resp.status_code == 200, (
        f"SS1.3: a cap breach is a per-file failure, HTTP stays 200 -- got "
        f"{resp.status_code}: {resp.text[:400]}"
    )
    entries = {e["name"]: e for e in resp.json()["documents"]}
    bad = entries[bad_name]
    assert bad["status"] == "failed", f"expected failed, got {bad}"
    assert bad["error"] == expected_error, (
        f"SS1.3 freezes this error string.\n  expected: {expected_error!r}\n"
        f"  actual:   {bad.get('error')!r}"
    )
    assert bad["id"] is None and bad["pages"] is None and bad["chunks"] == 0, bad
    assert bad.get("tables") == 0, f"SS1.3: failed entries report tables:0, got {bad.get('tables')!r}"
    assert entries[GOOD_TXT[0]]["status"] == "indexed", (
        "SS1.3: one file breaching a cap must not take the rest of the batch down: "
        f"{entries[GOOD_TXT[0]]}"
    )
    listed = f.client.get("/api/documents").json()
    names = [d["name"] for d in listed["documents"]]
    assert names == [GOOD_TXT[0]], f"SS1.3: nothing is persisted for a failed file, got {names}"
    uploads = f.storage / "uploads"
    on_disk = {p.name for p in uploads.iterdir()} if uploads.is_dir() else set()
    assert on_disk == {d["id"] for d in listed["documents"]}, (
        f"SS3.3: uploads/ contains exactly the manifest ids -- stray dirs: {on_disk}"
    )
    assert f.client.get("/api/health").status_code == 200, (
        "SS1.2: health is always 200 while the process is up -- a cap must never kill it"
    )
    return bad


# --------------------------------------------------------------------------
# SS1.3 / SS2 constants
# --------------------------------------------------------------------------
def test_allowed_exts_is_the_v12_tuple(fmt, qa):
    ingest = qa.backend_module("ingest")
    assert tuple(ingest.ALLOWED_EXTS) == ALLOWED_EXTS_V12, (
        f"SS2: ALLOWED_EXTS is frozen as {ALLOWED_EXTS_V12}, got {tuple(ingest.ALLOWED_EXTS)}"
    )


def test_unsupported_extension_message_lists_all_ten(fmt, qa):
    entry = upload_bytes(fmt.client, [("notes.exe", b"MZ binary")]).json()["documents"][0]
    assert entry["status"] == "failed"
    assert entry["error"] == qa.ERR.unsupported.format(ext=".exe"), (
        f"SS1.3 frozen string (it now names all ten extensions):\n"
        f"  expected: {qa.ERR.unsupported.format(ext='.exe')!r}\n  actual:   {entry['error']!r}"
    )


@pytest.mark.parametrize("name,value", sorted(EXTRACTION_CAPS.items()))
def test_extraction_cap_constants_are_frozen(fmt, qa, name, value):
    ingest = qa.backend_module("ingest")
    got = qa.require_attr(ingest, name, "SS2 ingest.py extraction caps")
    assert got == value, f"SS2 freezes {name} = {value}, got {got!r}"


def test_extraction_cap_exceeded_exception_exists(fmt, qa):
    ingest = qa.backend_module("ingest")
    exc = qa.require_attr(ingest, "ExtractionCapExceeded", "SS2 ingest.py")
    assert issubclass(exc, Exception)
    instance = exc("boom")
    assert getattr(instance, "message", None) == "boom", (
        "SS2: ExtractionCapExceeded carries attr `message` holding the exact SS1.3 string"
    )


# --------------------------------------------------------------------------
# .xlsx -- SS2 parsing table
# --------------------------------------------------------------------------
XLSX_SHEETS = {
    "Zephyr Q3 Metrics": [
        ["Metric", "Value"],
        ["Zephyr Labs revenue", "$77.3 million"],
        ["Zephyr Labs headcount", 1482],
        ["Zephyr Labs gross margin", "63%"],
    ],
    "Notes": [["Note"], ["Prepared by the Zephyr Labs finance team."]],
}


def test_xlsx_indexes_with_expected_manifest_fields(fmt):
    entry = index_one(fmt.client, "zephyr_metrics.xlsx", make_xlsx(XLSX_SHEETS))
    assert entry["pages"] is None, f"SS1.3: pages is null for every ext except .pdf/.pptx: {entry}"
    assert entry["chunks"] >= 1
    assert entry["tables"] == 2, (
        f"SS2: xlsx tables = one per NON-EMPTY worksheet (2 here), got {entry.get('tables')!r}"
    )


def test_xlsx_chunks_are_table_flagged_and_pageless(fmt):
    entry = index_one(fmt.client, "zephyr_metrics.xlsx", make_xlsx(XLSX_SHEETS))
    rows = previews(fmt.client, entry["id"])
    assert rows and all(r["has_table"] is True for r in rows), (
        f"SS2: has_table is true for EVERY xlsx block, got {[r['has_table'] for r in rows]}"
    )
    assert all(r["page"] is None for r in rows), [r["page"] for r in rows]


def test_xlsx_sheet_name_is_retrievable_text(fmt):
    entry = index_one(fmt.client, "zephyr_metrics.xlsx", make_xlsx(XLSX_SHEETS))
    text = all_text(fmt.client, entry["id"])
    assert "Zephyr Q3 Metrics" in text, (
        "SS2: each xlsx block starts with a `Sheet: {name}` header line so the sheet "
        f"name is retrievable. Parsed text head: {text[:300]!r}"
    )


def test_xlsx_numeric_cell_is_retrievable_through_a_real_query(fmt):
    """The anti-regression case: a real PROVIDER=none query, not a parser unit test."""
    entry = index_one(fmt.client, "zephyr_metrics.xlsx", make_xlsx(XLSX_SHEETS))
    resp = post_query(fmt.client, "What was the Zephyr Labs headcount?")
    assert resp.status_code == 200, resp.text[:400]
    body = resp.json()
    assert body["no_answer"] is False, (
        "a numeric cell that parsed but cannot be retrieved is a broken format: "
        f"{body['answer']!r}"
    )
    assert body["mode"] == "extractive" and body["model"] is None
    assert body["citations"], body
    haystack = body["answer"] + " " + " ".join(c["snippet"] for c in body["citations"])
    assert "1482" in haystack, (
        "SS2: openpyxl runs with data_only=True, so the numeric cell 1482 must appear "
        f"verbatim in the retrieved text. Got: {haystack[:400]!r}"
    )
    assert entry["id"] in {c["doc_id"] for c in body["citations"]}


def test_xlsx_currency_cell_is_retrievable_through_a_real_query(fmt):
    index_one(fmt.client, "zephyr_metrics.xlsx", make_xlsx(XLSX_SHEETS))
    body = post_query(fmt.client, "What was Zephyr Labs revenue?").json()
    assert body["no_answer"] is False, body["answer"]
    haystack = " ".join(c["snippet"] for c in body["citations"])
    assert "77.3" in haystack, f"revenue figure not retrievable from the sheet: {haystack[:300]!r}"


def test_xlsx_must_sniff_as_ooxml(fmt, qa):
    entry = upload_bytes(fmt.client, [("fake.xlsx", b"col,val\n1,2\n")]).json()["documents"][0]
    assert entry["status"] == "failed"
    assert entry["error"] == qa.ERR.sniff, (
        f"SS1.3: .xlsx must sniff as PK\\x03\\x04, got {entry['error']!r}"
    )


# --------------------------------------------------------------------------
# .pptx -- SS2 parsing table
# --------------------------------------------------------------------------
PPTX_SLIDES = [
    {"text": ["Orion Freight FY2027 Review", "Prepared for the board"]},
    {
        "text": ["Segment results"],
        "table": [["Segment", "Revenue"], ["Air", "$12.4M"], ["Sea", "$31.9M"]],
    },
    {"text": ["Outlook remains constructive for the Orion Freight network."]},
]


def test_pptx_pages_are_slide_numbers(fmt):
    entry = index_one(fmt.client, "orion_deck.pptx", make_pptx(PPTX_SLIDES))
    assert entry["pages"] == 3, (
        f"SS1.3: `pages` is the slide count for .pptx, got {entry['pages']!r}"
    )
    rows = previews(fmt.client, entry["id"])
    pages = sorted({r["page"] for r in rows})
    assert pages and all(isinstance(p, int) and p >= 1 for p in pages), (
        f"SS1.8: pptx chunks carry an int slide number >= 1, got {pages}"
    )
    assert max(pages) <= 3


def test_pptx_tables_counted_and_flagged_per_slide(fmt):
    entry = index_one(fmt.client, "orion_deck.pptx", make_pptx(PPTX_SLIDES))
    assert entry["tables"] == 1, (
        f"SS2: pptx tables = one per table SHAPE (1 here), got {entry.get('tables')!r}"
    )
    rows = previews(fmt.client, entry["id"])
    flagged = {r["page"] for r in rows if r["has_table"]}
    assert flagged == {2}, (
        "SS2: has_table is true only for slides that contained a table (slide 2), "
        f"got flagged slides {sorted(flagged)}"
    )


def test_pptx_table_cell_is_retrievable_through_a_real_query(fmt):
    index_one(fmt.client, "orion_deck.pptx", make_pptx(PPTX_SLIDES))
    body = post_query(fmt.client, "What was the Orion Freight Sea segment revenue?").json()
    assert body["no_answer"] is False, body["answer"]
    haystack = " ".join(c["snippet"] for c in body["citations"])
    assert "31.9" in haystack, f"pptx table cell not retrievable: {haystack[:300]!r}"


# --------------------------------------------------------------------------
# .html / .htm -- SS2: strip to inert text, tables not parsed (resolution 16)
# --------------------------------------------------------------------------
HTML_BODY = (
    "<h1>Atlas Mining &amp; Metals</h1>"
    "<p>FY2028 copper output reached 412 kilotonnes.</p>"
    "<table><tr><td>Copper</td><td>412kt</td></tr></table>"
)


@pytest.mark.parametrize("ext", [".html", ".htm"])
def test_html_strips_tags_scripts_styles_and_comments(fmt, ext):
    payload = make_html(body=HTML_BODY, script="var SECRETSCRIPT=1;", style=".x{color:SECRETSTYLE}")
    entry = index_one(fmt.client, f"atlas{ext}", payload)
    assert entry["pages"] is None, entry
    assert entry["tables"] == 0, (
        f"resolution 16: v1.2 does not parse HTML tables, tables must be 0, got {entry['tables']!r}"
    )
    text = all_text(fmt.client, entry["id"])
    for banned in ("SECRETSCRIPT", "SECRETSTYLE", "HIDDENCOMMENT", "<p>", "<table", "</h1>"):
        assert banned not in text, (
            f"SS2: script/style/comment content is dropped and tags are stripped -- "
            f"{banned!r} survived into the indexed text: {text[:300]!r}"
        )
    assert "Atlas Mining & Metals" in text, (
        f"SS2: entities are unescaped (&amp; -> &), got {text[:200]!r}"
    )
    assert "  " not in text, f"SS2: whitespace is collapsed: {text[:200]!r}"


def test_html_chunks_are_pageless_and_not_table_flagged(fmt):
    entry = index_one(fmt.client, "atlas.html", make_html(body=HTML_BODY))
    rows = previews(fmt.client, entry["id"])
    assert all(r["page"] is None for r in rows), [r["page"] for r in rows]
    assert all(r["has_table"] is False for r in rows), (
        "resolution 16: HTML tables are not parsed, so no chunk is table-flagged"
    )


def test_html_body_text_is_retrievable_through_a_real_query(fmt):
    index_one(fmt.client, "atlas.html", make_html(body=HTML_BODY))
    body = post_query(fmt.client, "What was Atlas Mining copper output?").json()
    assert body["no_answer"] is False, body["answer"]
    haystack = " ".join(c["snippet"] for c in body["citations"])
    assert "412" in haystack, f"html text not retrievable: {haystack[:300]!r}"


def test_html_with_no_text_is_a_clean_no_extractable_text(fmt, qa):
    payload = b"<!doctype html><html><head><style>p{}</style></head><body>   </body></html>"
    entry = upload_bytes(fmt.client, [("blank.html", payload)]).json()["documents"][0]
    assert entry["status"] == "failed"
    assert entry["error"] == qa.ERR.no_text, entry


def test_html_with_nul_bytes_fails_the_text_sniff(fmt, qa):
    entry = upload_bytes(fmt.client, [("bad.html", b"<p>ok</p>\x00<p>no</p>")]).json()["documents"][0]
    assert entry["status"] == "failed"
    assert entry["error"] == qa.ERR.sniff, (
        f"SS1.3: text sniff rejects NUL bytes, got {entry['error']!r}"
    )


# --------------------------------------------------------------------------
# .json -- SS2: key-path lines, iterative traversal
# --------------------------------------------------------------------------
JSON_DOC = {
    "company": "Vantage Grid",
    "metrics": [
        {"period": "FY2029", "revenue": "$88.6M", "customers": 2317},
        {"period": "FY2028", "revenue": "$71.2M", "customers": 1904},
    ],
    "notes": {"auditor": "Keller & Roth"},
}


def test_json_indexes_as_key_path_lines(fmt):
    entry = index_one(fmt.client, "vantage.json", make_json_bytes(JSON_DOC))
    assert entry["pages"] is None and entry["tables"] == 0, entry
    text = all_text(fmt.client, entry["id"])
    assert "metrics[0].revenue" in text, (
        f"SS2: json is flattened to `a.b[0].c: value` key-path lines, got {text[:400]!r}"
    )
    assert "notes.auditor" in text, text[:400]
    rows = previews(fmt.client, entry["id"])
    assert all(r["page"] is None and r["has_table"] is False for r in rows)


def test_json_leaf_value_is_retrievable_through_a_real_query(fmt):
    index_one(fmt.client, "vantage.json", make_json_bytes(JSON_DOC))
    body = post_query(fmt.client, "How many Vantage Grid customers were there?").json()
    assert body["no_answer"] is False, body["answer"]
    haystack = " ".join(c["snippet"] for c in body["citations"])
    assert "2317" in haystack, f"json leaf value not retrievable: {haystack[:300]!r}"


@pytest.mark.parametrize(
    "root", [[1, 2, 3], "bare string", 42, True], ids=["array", "string", "int", "bool"]
)
def test_json_accepts_any_root_type(fmt, root):
    resp = upload_bytes(fmt.client, [("root.json", make_json_bytes(root))])
    entry = resp.json()["documents"][0]
    assert entry["status"] in ("indexed", "failed"), entry
    if entry["status"] == "failed":
        assert entry["error"] == "no extractable text", (
            f"SS2: 'any root type allowed' -- a valid JSON root must not be a parse "
            f"failure, got {entry['error']!r}"
        )


def test_malformed_json_is_a_clean_parse_failure(fmt, qa):
    entry = upload_bytes(fmt.client, [("bad.json", b'{"a": 1,,}')]).json()["documents"][0]
    assert entry["status"] == "failed"
    assert entry["error"] == qa.ERR.parse, (
        f"SS2: a json parse failure is `failed to parse file`, got {entry['error']!r}"
    )


# --------------------------------------------------------------------------
# SS2 extraction caps -- clean per-file failure, never a 500, never a crash
# --------------------------------------------------------------------------
def test_json_depth_bomb_fails_cleanly(fmt, qa):
    """Real depth bomb: SS2 counts depth iteratively and must never recurse."""
    cap = EXTRACTION_CAPS["JSON_MAX_DEPTH"]
    payload = make_json_bytes(nested_json(cap + 8))
    assert_clean_per_file_failure(
        fmt, "deep.json", payload, qa.ERR.json_depth.format(cap=cap)
    )


def test_json_depth_bomb_does_not_blow_the_python_stack(fmt):
    """A recursive parser would raise RecursionError -> 500. SS2 mandates an explicit stack."""
    payload = make_json_bytes(nested_json(3000))
    resp = upload_bytes(fmt.client, [("verydeep.json", payload)])
    assert resp.status_code == 200, (
        f"SS1.3 (law): a depth bomb is a clean per-file failure, not a crash -- "
        f"HTTP {resp.status_code}: {resp.text[:300]}"
    )
    entry = resp.json()["documents"][0]
    assert entry["status"] == "failed" and entry["error"].startswith("json too deeply nested"), entry


def test_json_node_cap_fails_cleanly(fmt, qa):
    with patch_cap(qa, "JSON_MAX_NODES", 10):
        assert_clean_per_file_failure(
            fmt,
            "wide.json",
            make_json_bytes({f"k{i}": i for i in range(60)}),
            qa.ERR.json_nodes.format(cap=10),
        )


def test_xlsx_sheet_cap_fails_cleanly(fmt, qa):
    """Real cap breach: 51 sheets against the frozen XLSX_MAX_SHEETS = 50."""
    cap = EXTRACTION_CAPS["XLSX_MAX_SHEETS"]
    sheets = {f"S{i}": [["Metric", "Value"], ["rev", i]] for i in range(cap + 1)}
    assert_clean_per_file_failure(
        fmt, "toomany.xlsx", make_xlsx(sheets), qa.ERR.xlsx_sheets.format(cap=cap)
    )


def test_xlsx_cell_cap_fails_cleanly(fmt, qa):
    """Cap patched down (SS2 makes it a named module-top constant precisely for this)."""
    with patch_cap(qa, "XLSX_MAX_CELLS", 12):
        rows = [[f"c{c}" for c in range(6)] for _ in range(20)]  # 120 cells
        assert_clean_per_file_failure(
            fmt, "huge.xlsx", make_xlsx({"Big": rows}), qa.ERR.xlsx_cells.format(cap=12)
        )


def test_xlsx_cell_cap_message_reads_the_live_constant(fmt, qa):
    with patch_cap(qa, "XLSX_MAX_CELLS", 5):
        rows = [[f"c{c}" for c in range(5)] for _ in range(10)]
        entry = upload_bytes(fmt.client, [("huge2.xlsx", make_xlsx({"Big": rows}))]).json()
        error = entry["documents"][0].get("error", "")
    assert error == qa.ERR.xlsx_cells.format(cap=5), (
        "SS2: the cap constant must be read at call time so the SS1.3 message reports the "
        f"cap that was actually enforced. Got {error!r}"
    )


def test_pptx_slide_cap_fails_cleanly(fmt, qa):
    with patch_cap(qa, "PPTX_MAX_SLIDES", 2):
        deck = make_pptx([{"text": [f"Slide {i}"]} for i in range(5)])
        assert_clean_per_file_failure(
            fmt, "big.pptx", deck, qa.ERR.pptx_slides.format(cap=2)
        )


def test_html_size_cap_uses_the_folded_general_text_rail(fmt, qa):
    """r3 folded the html-specific cap into MAX_EXTRACTED_TEXT_CHARS at the SAME
    value, so HTML behavior is identical -- only the constant name and the frozen
    string change. The format-specific constant must be GONE, not merely unused:
    a stale second cap is how the two silently drift apart."""
    ingest = qa.backend_module("ingest")
    assert not hasattr(ingest, "HTML_MAX_TEXT_CHARS"), (
        "SS2 (r3): the html-specific text cap is removed, superseded by the "
        "format-agnostic rail"
    )
    with patch_cap(qa, "MAX_EXTRACTED_TEXT_CHARS", 200):
        payload = make_html(body="<p>" + ("Atlas output data. " * 200) + "</p>")
        assert_clean_per_file_failure(
            fmt, "big.html", payload, qa.ERR.text_cap.format(cap=200)
        )


# --------------------------------------------------------------------------
# SS2 ooxml_guard -- runs BEFORE any OOXML library opens the file
# --------------------------------------------------------------------------
BOMB_UNCOMPRESSED = 16 << 20  # 16 MB: well under OOXML_MAX_UNCOMPRESSED_BYTES (200 MB),
# so the guard is forced to trip on the RATIO check specifically, not the size check.
_BOMB_CACHE = {}


def _ratio_bomb():
    if "bytes" not in _BOMB_CACHE:
        _BOMB_CACHE["bytes"] = make_zip_bytes(
            [("[Content_Types].xml", b"<Types/>"), ("xl/big.bin", b"\0" * BOMB_UNCOMPRESSED)]
        )
    return _BOMB_CACHE["bytes"]


@pytest.mark.parametrize("name", ["bomb.xlsx", "bomb.docx", "bomb.pptx"])
def test_ooxml_ratio_bomb_fails_cleanly(fmt, qa, name):
    assert_clean_per_file_failure(fmt, name, _ratio_bomb(), qa.ERR.zip_bomb)


def test_ooxml_entry_count_bomb_fails_cleanly(fmt, qa):
    cap = EXTRACTION_CAPS["OOXML_MAX_ENTRIES"]
    entries = [(f"e{i}.xml", b"<x/>") for i in range(cap + 5)]
    assert_clean_per_file_failure(fmt, "many.xlsx", make_zip_bytes(entries), qa.ERR.zip_bomb)


def test_ooxml_guard_is_a_central_directory_read(fmt, qa, tmp_path):
    """SS2: ooxml_guard reads the central directory only -- no extraction, so it is fast
    and bounded even when the declared uncompressed size is enormous."""
    ingest = qa.backend_module("ingest")
    guard = qa.require_attr(ingest, "ooxml_guard", "SS2 ingest.py")
    exc = qa.require_attr(ingest, "ExtractionCapExceeded", "SS2 ingest.py")
    bomb = tmp_path / "bomb.xlsx"
    bomb.write_bytes(_ratio_bomb())
    with pytest.raises(exc) as caught:
        guard(bomb)
    assert caught.value.message == qa.ERR.zip_bomb, caught.value


def test_ooxml_guard_passes_a_normal_workbook(fmt, qa, tmp_path):
    ingest = qa.backend_module("ingest")
    guard = qa.require_attr(ingest, "ooxml_guard", "SS2 ingest.py")
    ok = tmp_path / "ok.xlsx"
    ok.write_bytes(make_xlsx(XLSX_SHEETS))
    guard(ok)  # must not raise


def test_zip_bomb_fixture_trips_ratio_but_not_size():
    """Sanity check on the fixture itself, not on the product."""
    raw = b"\0" * BOMB_UNCOMPRESSED
    ratio = len(raw) / max(1, len(zlib.compress(raw, 6)))
    assert ratio > EXTRACTION_CAPS["OOXML_MAX_COMPRESSION_RATIO"], ratio
    assert BOMB_UNCOMPRESSED < EXTRACTION_CAPS["OOXML_MAX_UNCOMPRESSED_BYTES"]
    assert len(_ratio_bomb()) < 1_000_000, "the bomb must stay tiny on the wire"


# --------------------------------------------------------------------------
# batch + store integrity across the new formats
# --------------------------------------------------------------------------
def test_mixed_format_batch_indexes_and_totals_add_up(fmt):
    items = [
        ("zephyr.xlsx", make_xlsx(XLSX_SHEETS)),
        ("orion.pptx", make_pptx(PPTX_SLIDES)),
        ("atlas.html", make_html(body=HTML_BODY)),
        ("vantage.json", make_json_bytes(JSON_DOC)),
    ]
    resp = upload_bytes(fmt.client, items)
    assert resp.status_code == 200, resp.text[:400]
    entries = resp.json()["documents"]
    assert [e["status"] for e in entries] == ["indexed"] * 4, [
        (e["name"], e.get("error")) for e in entries
    ]
    listed = fmt.client.get("/api/documents").json()
    totals = listed["totals"]
    assert totals["documents"] == 4
    assert totals["chunks"] == sum(e["chunks"] for e in entries)
    assert totals["tables"] == sum(e["tables"] for e in entries), (
        f"SS1.4: totals.tables sums the per-doc tables field, got {totals}"
    )
    assert totals["pages"] == sum(e["pages"] or 0 for e in entries), (
        f"SS1.4: totals.pages sums non-null pages (only the pptx has any), got {totals}"
    )
    for d in listed["documents"]:
        assert d["ext"] in ALLOWED_EXTS_V12, d


def test_duplicate_of_a_new_format_reports_the_existing_tables_count(fmt):
    payload = make_xlsx(XLSX_SHEETS)
    first = index_one(fmt.client, "zephyr.xlsx", payload)
    again = upload_bytes(fmt.client, [("zephyr_copy.xlsx", payload)]).json()["documents"][0]
    assert again["status"] == "duplicate", again
    assert again["id"] == first["id"] and again["chunks"] == first["chunks"]
    assert again["tables"] == first["tables"], (
        f"SS1.3: a duplicate reports the EXISTING doc's counts including tables, got {again}"
    )
    assert "error" not in again, f"SS1.3: duplicates carry no error field, got {again}"


def test_new_formats_survive_a_restart(tmp_path, samples, qa):
    """SS3.4(5): a store containing v1.2 extensions is not a corruption condition.

    Two SEQUENTIAL app builds (never nested -- app_client purges backend.app.* on
    both entry and exit, which is exactly how the harness simulates a restart).
    """
    storage = tmp_path / "restart-storage"
    with app_client(storage) as client:
        index_one(client, "vantage.json", make_json_bytes(JSON_DOC))
        index_one(client, "atlas.html", make_html(body=HTML_BODY))
        before = client.get("/api/documents").json()
    with app_client(storage) as client2:
        after = client2.get("/api/documents").json()
        assert after["totals"] == before["totals"], (
            f"reconciliation changed the store across a restart: {before['totals']} -> "
            f"{after['totals']}"
        )
        body = post_query(client2, "How many Vantage Grid customers were there?").json()
        assert "2317" in " ".join(c["snippet"] for c in body["citations"]), body


def test_json_bytes_helper_roundtrips():
    assert _json.loads(make_json_bytes(JSON_DOC).decode()) == JSON_DOC


# --------------------------------------------------------------------------
# SS2 (law) -- METADATA NEUTRALITY. Resolution 9 calls this "the highest-risk
# line in v1.2" for the 100% eval gate: BM25 tokenizes
# get_content(MetadataMode.EMBED), so any new metadata key left visible shifts
# the sparse token stream.
# --------------------------------------------------------------------------
V11_METADATA_KEYS = ["doc_id", "doc_name", "page", "chunk_ix"]


def _all_nodes(qa):
    return qa.backend_module("stores").get_store().nodes_for(None)


def test_v11_metadata_keys_keep_their_names_values_and_order(fmt, qa):
    index_one(fmt.client, "zephyr.xlsx", make_xlsx(XLSX_SHEETS))
    index_one(fmt.client, "orion.pptx", make_pptx(PPTX_SLIDES))
    for node in _all_nodes(qa):
        keys = list(node.metadata)
        assert keys[:4] == V11_METADATA_KEYS, (
            "SS2 (law): the four v1.1 metadata keys keep their exact names AND "
            f"insertion order; new keys append after them. Got {keys}"
        )
        assert isinstance(node.metadata["doc_id"], str)
        assert isinstance(node.metadata["chunk_ix"], int)


def test_every_post_v11_metadata_key_is_excluded_from_embed_and_llm(fmt, qa):
    index_one(fmt.client, "zephyr.xlsx", make_xlsx(XLSX_SHEETS))
    index_one(fmt.client, "orion.pptx", make_pptx(PPTX_SLIDES))
    index_one(fmt.client, "vantage.json", make_json_bytes(JSON_DOC))
    offenders = []
    for node in _all_nodes(qa):
        new_keys = set(node.metadata) - set(V11_METADATA_KEYS)
        embed_excl = set(node.excluded_embed_metadata_keys or [])
        llm_excl = set(node.excluded_llm_metadata_keys or [])
        for key in sorted(new_keys):
            if key not in embed_excl or key not in llm_excl:
                offenders.append(
                    f"{node.metadata['doc_name']} chunk {node.metadata['chunk_ix']}: "
                    f"{key!r} embed_excluded={key in embed_excl} llm_excluded={key in llm_excl}"
                )
    assert not offenders, (
        "SS2 (law) / resolution 9: EVERY metadata key added in v1.2 or later "
        "(has_table, sheet, slide, ...) MUST appear in both "
        "excluded_embed_metadata_keys and excluded_llm_metadata_keys. A visible key "
        "shifts the BM25 token stream and puts the 100% eval gate at risk.\n  "
        + "\n  ".join(offenders)
    )


def test_bm25_token_stream_shows_only_the_chunk_text(fmt, qa):
    """The direct consequence of the law above, asserted on the actual EMBED render."""
    from llama_index.core.schema import MetadataMode

    index_one(fmt.client, "zephyr.xlsx", make_xlsx(XLSX_SHEETS))
    index_one(fmt.client, "orion.pptx", make_pptx(PPTX_SLIDES))
    for node in _all_nodes(qa):
        embed_view = node.get_content(metadata_mode=MetadataMode.EMBED)
        text = node.get_content()
        assert embed_view.endswith(text), (
            f"the EMBED render must end with the chunk text:\n{embed_view[:300]!r}"
        )
        header = embed_view[: len(embed_view) - len(text)]
        rendered_keys = [
            line.split(":", 1)[0] for line in header.splitlines() if ":" in line
        ]
        assert rendered_keys == V11_METADATA_KEYS, (
            "SS2 (law) / resolution 9: the BM25 token stream renders EXACTLY the four "
            "v1.1 metadata keys, in order, and nothing else. Any v1.2 key that appears "
            f"here shifts the sparse tokens and endangers the 100% eval gate. Got "
            f"{rendered_keys}"
        )
        for key in set(node.metadata) - set(V11_METADATA_KEYS):
            assert f"{key}:" not in header, f"{key!r} leaked into the BM25 stream"


# --------------------------------------------------------------------------
# SS3.2 -- `tables` in the manifest, and v1.1 backward compatibility (law)
# --------------------------------------------------------------------------
def test_manifest_entries_carry_tables_and_table_total_agrees(fmt, qa):
    index_one(fmt.client, "zephyr.xlsx", make_xlsx(XLSX_SHEETS))
    index_one(fmt.client, "atlas.html", make_html(body=HTML_BODY))
    store = qa.backend_module("stores").get_store()
    manifest = store.get_manifest()
    assert manifest, "manifest is empty"
    for entry in manifest:
        assert isinstance(entry.get("tables"), int) and entry["tables"] >= 0, (
            f"SS3.2: every entry written from v1.2 onward must include tables:int>=0, got {entry}"
        )
    total = qa.require_attr(store, "table_total", "SS2 stores.py")()
    assert total == sum(e["tables"] for e in manifest), (
        f"SS2: table_total() is the manifest sum, got {total}"
    )
    assert total == fmt.client.get("/api/documents").json()["totals"]["tables"]


def test_a_v11_manifest_without_tables_boots_unchanged(tmp_path, samples, qa):
    """SS3.2 (law) + SS3.4(5): a v1.1 entry has NO `tables` key -- not `tables: 0`.
    It reads as 0 everywhere, it is NOT corruption, and startup performs NO
    migration and NO manifest rewrite; rewriting a healthy store at boot is the
    banned 'silent rebuild'.

    This is also the exact shape that can make a CORRECT v1.2 system look broken:
    a legacy store reports `tables: 0` / `has_table: false` truthfully, and a
    re-upload of the same bytes hits sha256 dedupe (SS1.3) and hands back that
    legacy entry rather than re-indexing it. Both behaviours are pinned below so
    nobody "fixes" them into a self-healing rewrite.
    """
    import json
    import os
    import shutil

    storage = tmp_path / "storage"
    payloads = [
        ("vantage.json", make_json_bytes(JSON_DOC)),
        ("atlas.html", make_html(body=HTML_BODY)),
        ("rows.csv", b"col_a,col_b\n1,2\n3,4\n"),  # a doc that really does have a table
    ]
    with app_client(storage) as client:
        for name, data in payloads:
            index_one(client, name, data)
        modern = client.get("/api/documents").json()
    assert modern["totals"]["tables"] >= 1, (
        f"fixture sanity: the v1.2 store must have a non-zero tables total first, "
        f"otherwise 'reads as 0' proves nothing. Got {modern['totals']}"
    )

    # --- age an ISOLATED COPY down to a genuine v1.1 shape -----------------
    # The copy matters: the artifact under test is then a pristine file tree that
    # nothing else in the process holds a handle to, so this test cannot be
    # perturbed by (nor perturb) the store the first app just closed.
    #
    # The ageing rule is DERIVED FROM THE DATA, not from a hardcoded list of v1.2
    # key names: "v1.1 shape" is defined by SS2's metadata-neutrality law as exactly
    # the four v1.1 keys and nothing else. That way a new metadata key added in a
    # later version is aged away automatically instead of silently surviving.
    legacy = tmp_path / "legacy-storage"
    shutil.copytree(storage, legacy)

    manifest_path = legacy / "manifest.json"
    data = json.loads(manifest_path.read_text(encoding="utf-8"))
    for entry in data["documents"]:
        assert "tables" in entry, f"v1.2 must write the key in the first place: {entry}"
        entry.pop("tables")
    manifest_path.write_text(json.dumps(data), encoding="utf-8")

    docstore_path = legacy / "docstore.json"
    exclusion_lists = ("excluded_embed_metadata_keys", "excluded_llm_metadata_keys")

    def _age_to_v11(node):
        """Reduce every node's metadata to the four v1.1 keys and empty the
        exclusion lists, which only exist because post-v1.1 keys do."""
        if isinstance(node, dict):
            md = node.get("metadata")
            if isinstance(md, dict):
                for key in [k for k in md if k not in V11_METADATA_KEYS]:
                    md.pop(key)
            for name in exclusion_lists:
                if isinstance(node.get(name), list):
                    node[name] = []
            for value in node.values():
                _age_to_v11(value)
        elif isinstance(node, list):
            for value in node:
                _age_to_v11(value)

    def _post_v11_paths(node, path="$"):
        hits = []
        if isinstance(node, dict):
            md = node.get("metadata")
            if isinstance(md, dict):
                extra = sorted(k for k in md if k not in V11_METADATA_KEYS)
                if extra:
                    hits.append(f"{path}.metadata has {extra}")
            for name in exclusion_lists:
                if node.get(name):
                    hits.append(f"{path}.{name} = {node[name]!r}")
            for key, value in node.items():
                hits += _post_v11_paths(value, f"{path}.{key}")
        elif isinstance(node, list):
            for i, value in enumerate(node):
                hits += _post_v11_paths(value, f"{path}[{i}]")
        return hits

    parsed = json.loads(docstore_path.read_text(encoding="utf-8"))
    assert _post_v11_paths(parsed), (
        "fixture sanity: the v1.2 store must carry post-v1.1 node metadata before "
        "ageing, otherwise 'reads as 0/false' proves nothing"
    )
    _age_to_v11(parsed)
    docstore_path.write_text(json.dumps(parsed), encoding="utf-8")

    leftovers = _post_v11_paths(json.loads(docstore_path.read_text(encoding="utf-8")))
    assert not leftovers, (
        "the aged docstore must carry nothing but the four v1.1 metadata keys -- "
        f"otherwise this test is not exercising a real pre-v1.2 store.\n  {leftovers[:8]}"
    )

    before = {p: (p.read_bytes(), os.stat(p).st_mtime_ns) for p in (manifest_path, docstore_path)}

    # --- a v1.2 process must serve that store as-is ------------------------
    with app_client(legacy) as client:
        assert client.get("/api/health").status_code == 200, (
            "SS3.4(5): a manifest entry missing `tables` is explicitly NOT corruption"
        )
        listed = client.get("/api/documents").json()
        assert all(d["tables"] == 0 for d in listed["documents"]), (
            f"SS3.2 (law): a missing tables key reads as 0 (`d.get('tables', 0) or 0`) "
            f"everywhere -- never null, never absent from the response. Got "
            f"{[(d['name'], d.get('tables')) for d in listed['documents']]}"
        )
        assert listed["totals"]["tables"] == 0, listed["totals"]

        for doc in listed["documents"]:
            rows = client.get(f"/api/documents/{doc['id']}/chunks").json()["chunks"]
            assert rows, doc
            assert all(r["has_table"] is False for r in rows), (
                f"SS1.8: chunks of pre-v1.2 documents carry no has_table metadata and "
                f"report False -- absence is not corruption. {doc['name']}: {rows}"
            )
            assert all(isinstance(r["has_table"], bool) for r in rows), rows

        # The dedupe path: re-uploading identical bytes returns the LEGACY entry.
        for name, payload in payloads:
            again = upload_bytes(client, [(name, payload)]).json()["documents"][0]
            assert again["status"] == "duplicate", (
                f"SS1.3: identical bytes are a sha256 duplicate, got {again}"
            )
            assert again["tables"] == 0, (
                "SS1.3: a duplicate reports the EXISTING doc's counts -- for a legacy "
                f"entry that is 0, and re-indexing it would be the banned rewrite. Got {again}"
            )
            assert "error" not in again, again

        body = post_query(client, "How many Vantage Grid customers were there?").json()
        assert "2317" in " ".join(c["snippet"] for c in body["citations"]), body

    for path, (raw_bytes, mtime) in before.items():
        assert path.read_bytes() == raw_bytes, (
            f"SS3.2 (law): startup performs NO migration and NO rewrite -- {path.name} "
            "must be byte-identical after boot, dedupe and query"
        )
        assert os.stat(path).st_mtime_ns == mtime, f"{path.name} was rewritten in place"


# --------------------------------------------------------------------------
# SS1.3 / SS2 -- `tables` and `has_table` for EVERY extension.
#
# Why this section exists: SS2 says "existing formats are unchanged", which reads
# as though the pdf/docx paths need no new coverage. But the COUNT PLUMBING is new
# in v1.2 even where the parser is not, and an unchanged parser feeding a new field
# is exactly where that field silently stays 0. The synthetic .xlsx/.pptx fixtures
# above carry tables and so covered their own paths; nothing exercised the
# pre-existing pdf/docx paths, and `tables` was 0 for every document in a live
# re-index while every shape assertion stayed green.
#
# The lesson generalized: assert the CONTRACT-DERIVED VALUE, never merely presence
# or type. A field that is always its default is invisible to both.
# --------------------------------------------------------------------------
def _tables_ground_truth(path, ext):
    """The SS2 rule computed independently, with the library SS2 names."""
    if ext == ".pdf":
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            return sum(len(page.extract_tables() or []) for page in pdf.pages)
    if ext == ".docx":
        import docx

        return len(docx.Document(str(path)).tables)
    raise AssertionError(f"no ground-truth rule for {ext}")


def _pdf_page_count(path):
    import pypdf

    return len(pypdf.PdfReader(str(path)).pages)


# key -> builder(samples) -> (filename, bytes, expected_tables, SS2 rule text)
TABLE_RULE_CASES = {
    "pdf": lambda s: (
        s["meridian"].name, s["meridian"].read_bytes(),
        _tables_ground_truth(s["meridian"], ".pdf"), "number of pdfplumber tables",
    ),
    "docx": lambda s: (
        s["helios"].name, s["helios"].read_bytes(),
        _tables_ground_truth(s["helios"], ".docx"), "len(document.tables)",
    ),
    "txt": lambda s: (s["northwind"].name, s["northwind"].read_bytes(), 0, "TXT/MD => 0"),
    "md": lambda s: ("notes.md", b"# Title\n\nProse with the figure 42 in it.\n", 0, "TXT/MD => 0"),
    "csv-with-rows": lambda s: (
        "rows.csv", b"col_a,col_b\n1,2\n3,4\n", 1, "CSV => 1 when the file has >=1 data row",
    ),
    "csv-header-only": lambda s: (
        "hdr.csv", b"col_a,col_b\n", 0, "CSV => 0 when there is no data row",
    ),
    "xlsx-one-sheet": lambda s: (
        "one.xlsx", make_xlsx({"One": [["M", "V"], ["r", 1]]}), 1, "one per non-empty worksheet",
    ),
    "xlsx-two-sheets": lambda s: (
        "two.xlsx",
        make_xlsx({"One": [["M", "V"], ["r", 1]], "Two": [["M", "V"], ["s", 2]]}),
        2, "one per non-empty worksheet",
    ),
    "pptx-no-table": lambda s: (
        "plain.pptx", make_pptx([{"text": ["only prose on this slide"]}]),
        0, "one per table shape found",
    ),
    "pptx-two-tables": lambda s: (
        "two.pptx",
        make_pptx([
            {"text": ["a"], "table": [["h", "i"], ["1", "2"]]},
            {"text": ["b"], "table": [["j", "k"], ["3", "4"]]},
        ]),
        2, "one per table shape found",
    ),
    "html": lambda s: (
        "page.html", make_html(body="<p>prose</p><table><tr><td>x</td></tr></table>"),
        0, "resolution 16: v1.2 does not parse HTML tables",
    ),
    "htm": lambda s: (
        "page.htm", make_html(body="<p>prose</p><table><tr><td>x</td></tr></table>"),
        0, "resolution 16: v1.2 does not parse HTML tables",
    ),
    "json": lambda s: ("doc.json", make_json_bytes(JSON_DOC), 0, "JSON => 0"),
}


@pytest.mark.parametrize("case", sorted(TABLE_RULE_CASES))
def test_tables_matches_the_contract_rule_for_every_extension(fmt, samples, case):
    name, data, expected, rule = TABLE_RULE_CASES[case](samples)
    entry = index_one(fmt.client, name, data)
    assert entry["tables"] == expected, (
        f"SS2 rule for {name}: {rule} => {expected}. The upload response reported "
        f"{entry['tables']!r}. (A field that is always 0 passes every presence and type "
        "check ever written -- assert the derived VALUE.)"
    )
    listed = fmt.client.get("/api/documents").json()
    assert listed["documents"][0]["tables"] == expected, (
        f"SS1.4/SS3.3 (law): the list response reads `tables` from the manifest and must "
        f"agree with the upload response. Got {listed['documents'][0]['tables']!r}"
    )
    assert listed["totals"]["tables"] == expected, listed["totals"]


@pytest.mark.parametrize("case", sorted(TABLE_RULE_CASES))
def test_has_table_agrees_with_the_documents_table_count(fmt, samples, case):
    """SS1.8/SS2 (law): has_table is inherited PER BLOCK, so a document whose
    `tables` is non-zero must flag at least one chunk, and one whose `tables` is 0
    must flag none. Same always-default trap, one level down."""
    name, data, expected, rule = TABLE_RULE_CASES[case](samples)
    entry = index_one(fmt.client, name, data)
    rows = previews(fmt.client, entry["id"])
    flagged = [r["chunk_ix"] for r in rows if r["has_table"]]
    if expected:
        assert flagged, (
            f"{name}: tables={expected} ({rule}) but NO chunk carries has_table -- the "
            f"per-block inheritance is dropped. Chunks: {rows}"
        )
    else:
        assert not flagged, (
            f"{name}: tables=0 ({rule}) yet chunks {flagged} are flagged has_table"
        )


def test_committed_sample_pdf_and_docx_report_their_real_table_counts(fmt, samples):
    """The regression this section was written for, on the REAL committed samples
    rather than synthetic fixtures -- ground truth computed here with pdfplumber and
    python-docx, the two libraries SS2 names for these formats."""
    for key, ext in (("meridian", ".pdf"), ("helios", ".docx")):
        path = samples[key]
        expected = _tables_ground_truth(path, ext)
        assert expected >= 1, (
            f"fixture sanity: {path.name} must actually contain a table for this test to "
            f"mean anything (ground truth computed {expected})"
        )
        entry = index_one(fmt.client, path.name, path.read_bytes())
        assert entry["tables"] == expected, (
            f"SS2: {path.name} tables must equal the parser's own count ({expected}), got "
            f"{entry['tables']!r}"
        )
        rows = previews(fmt.client, entry["id"])
        assert any(r["has_table"] for r in rows), (
            f"{path.name} contains {expected} table(s) but no chunk is flagged: {rows}"
        )


def test_committed_sample_pdf_reports_its_real_page_count(fmt, samples):
    """`pages` is the other per-doc int that a broken path leaves at a default."""
    path = samples["meridian"]
    expected = _pdf_page_count(path)
    assert expected >= 2, f"fixture sanity: the sample PDF should be multi-page, got {expected}"
    entry = index_one(fmt.client, path.name, path.read_bytes())
    assert entry["pages"] == expected, (
        f"SS1.3: `pages` is the pypdf page count for .pdf ({expected}), got {entry['pages']!r}"
    )
    assert fmt.client.get("/api/documents").json()["totals"]["pages"] == expected


def test_chunk_ix_actually_advances_on_a_multi_chunk_document(fmt, samples):
    """`chunk_ix` contiguous-from-0 is trivially satisfied by a single-chunk doc;
    this pins that the counter really increments."""
    path = samples["meridian"]
    entry = index_one(fmt.client, path.name, path.read_bytes())
    rows = previews(fmt.client, entry["id"])
    assert len(rows) >= 2, f"fixture sanity: expected a multi-chunk PDF, got {len(rows)}"
    assert max(r["chunk_ix"] for r in rows) == len(rows) - 1 >= 1, [r["chunk_ix"] for r in rows]
    assert len({r["preview"] for r in rows}) == len(rows), (
        "distinct chunks must have distinct previews -- identical previews mean the head "
        "is being taken from the wrong text"
    )
